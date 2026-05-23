from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "source"))

from simulation import Scenario, monte_carlo_patrol, monte_carlo_queue, recognition_metrics  # noqa: E402


TITLE = "大型活动安防中的随机到达、排队拥堵与巡逻发现：高阶叙事性模拟"
DOCX_NAME = "midterm_report.docx"


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text="", style=None, align=None, first_line=True):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, name="黑体")
    else:
        set_run_font(run, size=13, bold=True, name="黑体")
    return paragraph


def add_formula(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(f"<eq>{text}</eq>")
    set_run_font(run, size=12, name="Cambria Math")
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, name="宋体")
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=10.5, bold=True, name="黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.2
                for run in paragraph.runs:
                    set_run_font(run, size=10.5)
    doc.add_paragraph()
    return table


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def pct(value):
    return f"{value * 100:.1f}%"


def main():
    figures = ROOT / "report" / "figures"
    output = ROOT / "report" / DOCX_NAME
    metrics = json.loads((figures / "metrics.json").read_text(encoding="utf-8"))

    scenario = Scenario()
    cases = [
        ("基准方案", scenario),
        ("错峰入场", Scenario(peak_concentration=0.82)),
        ("增开通道", Scenario(gates=44)),
        ("组合优化", Scenario(peak_concentration=0.82, gates=44, service_rate=21.0, patrols=8, smart_bias=0.85, false_alarm=0.0022)),
    ]
    case_rows = []
    for name, case in cases:
        q = monte_carlo_queue(case, trials=300, seed=202405)
        patrol = monte_carlo_patrol(case.patrols, case.smart_bias, trials=300, seed=202405)
        rec = recognition_metrics(case)
        case_rows.append(
            [
                name,
                case.gates,
                f"{case.peak_concentration:.2f}",
                pct(q["risk_probability"]),
                f"{q['max_queue_mean']:.0f}",
                f"{q['avg_wait_mean']:.2f}",
                f"{patrol['median_hit_time']:.0f}",
                pct(rec["posterior_true_given_alert"]),
                f"{rec['review_police_hours']:.0f}",
            ]
        )

    doc = Document()
    style_document(doc)

    # ── 封面 ────────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    cover = [
        ("2026 年《概率论与随机过程》小组作业", 18, True),
        ("", 12, False),
        (f"题目：{TITLE}", 18, True),
        ("", 12, False),
        ("学院：国际学院", 14, False),
        ("年级：2024级", 14, False),
        ("小组编号：第 8 组", 14, False),
        ("成员：2024213209 王方宇（模型推导与报告撰写）", 14, False),
        ("      2024213181 黄骏杰（交互网页与仿真实现）", 14, False),
        ("      2024213211 刘益泽（图表分析与资料整理）", 14, False),
        ("      2024213201 李克轩（展示设计与复核）", 14, False),
        ("指导教师：李丽华", 14, False),
        ("二〇二六年五月", 14, False),
    ]
    for text, size, bold in cover:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, name="黑体" if bold else "宋体")

    doc.add_page_break()

    # ── 摘要 ────────────────────────────────────────────────────────────────
    add_heading(doc, "摘 要", 1)
    add_paragraph(
        doc,
        "大型演唱会等密集性活动面临入场拥堵、报警误判和巡逻响应三类核心安防风险。"
        "本文以概率论与随机过程为工具，围绕\u201c随机到达\u2014门区排队\u2014报警可信度\u2014场内巡逻\u201d构建了一套高阶叙事性仿真框架，"
        "并配套两套交互系统（纯前端叙事版与 React 数据仪表板版）加以可视化呈现。"
        "观众入场过程采用非齐次泊松过程刻画双峰波形，安检口服务以离散 M/M/c 排队近似，"
        "报警可信度用贝叶斯后验概率揭示低基率效应，场内巡逻则转化为二维随机游走的首达时问题。"
        f"基准场景（36 通道、集中度 1.20）下，300 次蒙特卡洛模拟显示最大队列超过 2500 人阈值的概率为 "
        f"{pct(metrics['queue']['risk_probability'])}，平均最大队列约 {metrics['queue']['max_queue_mean']:.0f} 人，"
        f"90 分位最大队列达到 {metrics['queue']['max_queue_p90']:.0f} 人；"
        f"识别系统在 95% 灵敏度与 0.5% 误报率下，报警后验概率仅约 {pct(metrics['recognition']['posterior_true_given_alert'])}，"
        f"单场误报核查耗费约 {metrics['recognition']['review_police_hours']:.0f} 警员小时。"
        "结果表明，单次仿真的平均值会系统性地低估尾部风险；"
        "错峰入场、通道冗余、降低误报与热点巡逻的组合策略在四个维度上均优于单一措施，"
        "为大型活动安防的量化决策提供了可复现的参考依据。",
    )
    add_paragraph(doc, "关键词：非齐次泊松过程；M/M/c 排队论；贝叶斯后验概率；随机游走首达时；蒙特卡洛模拟；大型活动安防", first_line=False)

    # ── 目录 ────────────────────────────────────────────────────────────────
    add_heading(doc, "目录", 1)
    toc_items = [
        "第一章 绪论",
        "第二章 相关理论与模型构建",
        "第三章 叙事性模拟设计与交互实现",
        "第四章 模拟结果与概率解释",
        "第五章 安防策略建议与方案比较",
        "第六章 学习反思报告",
        "参考文献",
        "附录（代码与运行说明）",
    ]
    for item in toc_items:
        add_paragraph(doc, item, first_line=False)
    doc.add_page_break()

    # ── 第一章 绪论 ─────────────────────────────────────────────────────────
    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_paragraph(
        doc,
        "近年来，大型演唱会、体育赛事和节庆活动的规模持续扩大，单场观众数量动辄超过数万人。"
        "安防压力通常在开场前 30 至 60 分钟内骤然集中，形成入场波峰。"
        "实际风险并不只来自总人数，而来自人群到达的随机波动、安检服务能力的不确定性、"
        "识别报警系统的低基率误判，以及巡逻队伍在二维空间中的搜索效率。"
        "若只使用平均人数或静态图表描述，很容易忽略队列突然积累、误报大量占用警力和异常点迟迟未被发现等尾部风险。",
    )
    add_heading(doc, "1.2 探究目标", 2)
    add_paragraph(
        doc,
        "本项目选择\u201c大型活动安防\u201d作为小切口，目标不是复刻真实场馆的全部复杂性，"
        "而是用概率论与随机过程中的典型工具讲清一个完整决策故事：观众如何随机到达，"
        "队列如何积累并形成拥堵风险，报警信号如何在低基率下被解读，巡逻如何以随机游走方式搜索异常，"
        "以及如何从参数对比中提炼安防优化策略。"
        "交互作品对应作业说明中\u201c高阶：叙事性模拟\u201d的要求，将理论推导、动态仿真、关键指标和可调控件"
        "整合为两套 Web 应用，形成可探索的量化分析报告。",
    )

    # ── 第二章 模型构建 ─────────────────────────────────────────────────────
    add_heading(doc, "第二章 相关理论与模型构建", 1)
    
    add_heading(doc, "2.1 非齐次泊松到达过程", 2)
    add_paragraph(
        doc,
        "在实际大型活动中，观众到达大门的时间分布并不是均匀的，而是随着开场时间的临近呈现出强烈的波动性与时间相关性。"
        "为了精确刻画这一非平稳随机过程，本项目采用非齐次泊松过程（Non-homogeneous Poisson Process, NHPP）进行建模。"
        "设 A(t) 为从开门（t=0）到第 t 分钟累计到达的观众总数。由于到达率随时间变化，我们定义一个随时间变化的时变强度函数 λ(t) ≥ 0。"
        "根据非齐次泊松过程的数学定义，在任意时间区间 (t_1, t_2] 内，到达人数增量服从泊松分布，其概率质量函数为：",
    )
    add_formula(doc, "P(N(t_2) - N(t_1) = k) = ((Λ(t_1, t_2))^k · e^(-Λ(t_1, t_2)))/(k!)   (k = 0, 1, 2, ...)")
    add_paragraph(
        doc,
        "其中，累计到达人数的期望值（即均值积分）为：",
    )
    add_formula(doc, "Λ(t_1, t_2) = ∫_(t_1)^(t_2) λ(u) du")
    add_paragraph(
        doc,
        "且在不重叠的时间区间内，新到达的人数增量相互独立。在本项目中，为了模拟演唱会前观众的到达特征（包含早到峰和临场峰），"
        "我们将强度函数 λ(t) 构造为背景常数与两个高斯钟形分布（双峰分布）的加权叠加：",
    )
    add_formula(doc, "λ(t) = C · [ λ_0 + γ · ( e^(-0.5 · ((t-μ_1)/σ_1)^2) + 0.72 · e^(-0.5 · ((t-μ_2)/σ_2)^2) ) ]")
    add_paragraph(
        doc,
        "其中背景强度 λ₀ = 0.22，到达集中度为 γ，早到峰均值 μ₁ = 68 分钟，标准差 σ₁ = 18 分钟；"
        "临场峰均值 μ₂ = 102 分钟，标准差 σ₂ = 9 分钟。归一化系数 C 用于确保在 t = 120 分钟时总观众数期望值为 60000，"
        "即积分约束条件。系统模拟时，采用 Lewis & Shedler (1979) 的接受-拒绝减薄算法（Thinning Algorithm）生成随机到达路径，"
        "从而精确复现非平稳泊松增量的物理性质。",
    )

    add_heading(doc, "2.2 门区排队模型（离散 M/M/c 近似）", 2)
    add_paragraph(
        doc,
        "安检区设置 c 个并行通道，单通道平均服务率为 μ 人/分钟。如果使用经典的连续时间 M_t/M/c 排队模型，"
        "描述系统瞬时队长概率分布 P_n(t) = P(Q(t) = n) 的是一组无穷维一阶线性微分方程组（即 Kolmogorov 向前方程）：",
    )
    add_formula(doc, "(dP_0(t))/(dt) = -λ(t)P_0(t) + μ P_1(t)")
    add_formula(doc, "(dP_n(t))/(dt) = λ(t)P_(n-1)(t) - (λ(t) + n μ)P_n(t) + (n+1)μ P_(n+1)(t)   (1 ≤ n < c)")
    add_formula(doc, "(dP_n(t))/(dt) = λ(t)P_(n-1)(t) - (λ(t) + c μ)P_n(t) + c μ P_(n+1)(t)   (n ≥ c)")
    add_paragraph(
        doc,
        "该系统在到达率 λ(t) 随时间变化时极难解析求解。为了计算可行性并支持实时队长演化，我们将系统在时间轴上离散化（步长 Δt = 1 分钟）。"
        "设第 t 分钟到达数为 A(t) ~ Poisson(λ(t)Δt)，系统在该分钟的最大服务潜力为 S(t) ~ Poisson(cμΔt)。"
        "则队长的迭代关系为：",
    )
    add_formula(doc, "Q(t+1) = max(0, Q(t) + A(t) - S(t))")
    add_paragraph(
        doc,
        "该公式描述了一个在 0 处具有反射边界的非齐次一维随机游走。为了度量系统风险，我们通过 300 次蒙特卡洛独立模拟"
        "来计算尾部拥堵风险概率 P(max Q(t) > L)，以作为安全管理决策的依据。",
    )

    add_heading(doc, "2.3 报警可信度的贝叶斯解释", 2)
    add_paragraph(
        doc,
        "设 T 表示“该个体为真实风险人员”，B 表示“系统发出报警”。"
        "在低基率场景下，即使灵敏度较高，报警的后验概率也远低于直觉。应用贝叶斯公式得：",
    )
    add_formula(doc, "P(T|B) = (P(B|T) · P(T))/(P(B|T) · P(T) + P(B|T̄) · P(T̄))")
    add_paragraph(
        doc,
        "若总观众数 N = 60000，风险个体数 N_T = 50，则先验概率 P(T) = 50/60000 ≈ 0.0833%。"
        "系统灵敏度 α = P(B|T) = 95%，误报率 β = P(B|T̄) = 0.5%。代入计算：",
    )
    add_formula(doc, "P(T|B) = (0.95 × 0.000833)/(0.95 × 0.000833 + 0.005 × 0.999167) ≈ 13.7%")
    add_paragraph(
        doc,
        "这意味着报警中绝大多数是误报。设每次误报需派遣 k = 2 名警员，核查耗时 t_check = 20 分钟（即 1/3 小时），"
        "则整场活动因误报带来的警力核查成本（警员小时）计算为：",
    )
    add_formula(doc, "H_(review) = E[N_(false)] × k × (t_check/60) = (N - N_T) · β × 2 × (20/60) ≈ 199.8 警员小时")
    add_paragraph(
        doc,
        "公式解释了低基率效应下警力浪费的根本原因，并证明降低误报率的边际收益远高于提升灵敏度。",
    )

    add_heading(doc, "2.4 巡逻搜索与随机游走首达时", 2)
    add_paragraph(
        doc,
        "将场馆内部抽象为 W × H （12×8）的二维网格，N_P 支独立的巡逻队初始位置分布在入口和角落。"
        "设目标异常事件位于网格 x* = (8, 5)。首次被任一巡逻小组发现的时间定义为首达时随机变量 τ：",
    )
    add_formula(doc, "τ = inf { t ≥ 0 : ∃ i ∈ {1, 2, ..., N_P}, X_i(t) = x* }")
    add_paragraph(
        doc,
        "每支队伍以偏向度 p （曼哈顿距离减小方向）或以 1-p 随机进行网格移动，其位置状态转移满足马尔可夫链性质。"
        "单支队伍的期望首达时间 v(x) = E[τ | X_0 = x] 满足如下离散差分方程系统（泊松方程）：",
    )
    add_formula(doc, "v(x*) = 0,   v(x) = 1 + ∑_(y ∈ S) P(x, y) v(y)   (∀ x ≠ x*)")
    add_paragraph(
        doc,
        "多支队伍独立搜索时，相当于多个独立过程的竞争危险率。偏向度 p 可以改变转移矩阵的谱性质，"
        "不仅压缩了首达时间的中位数，也大幅缩短了其 90% 分位的右尾风险。",
    )

    # ── 第三章 交互实现 ─────────────────────────────────────────────────────
    add_heading(doc, "第三章 叙事性模拟设计与交互实现", 1)
    add_heading(doc, "3.1 故事线设计", 2)
    add_paragraph(
        doc,
        "交互作品按五幕叙事结构展开，每一幕对应一个核心概率概念，避免将图表堆叠为静态展示。"
        "第一幕：展示非均匀到达的人流双峰——通过实时动态曲线让用户感受到'随机到达并非匀速'；"
        "第二幕：展示服务能力不足时队列如何逐步积累，强调单次路径可能掩盖尾部风险；"
        "第三幕：用贝叶斯后验概率解释为何即使识别准确率较高，报警仍以误报为主；"
        "第四幕：将场内巡逻可视化为在网格上移动的随机游走路径，展示热点偏向的搜索加速效果；"
        "第五幕：以并排表格和参数对比形式综合呈现四种安防方案的量化差异，形成从风险识别到策略选择的完整闭环。",
    )

    add_heading(doc, "3.2 参数设置", 2)
    add_paragraph(
        doc,
        "下表列出仿真中所有可调参数的基准值及其含义，"
        "用户可在交互界面中实时修改这些参数并观察各指标的响应变化。",
        first_line=True,
    )
    add_table(
        doc,
        ["参数", "基准值", "含义"],
        [
            ["总观众数", "60 000", "进入场馆的总人数规模"],
            ["开放时长", "120 分钟", "开场前两小时开始入场"],
            ["安检通道", "36 个", "并行服务台数量（c）"],
            ["单通道效率", "20 人/分钟", "平均服务能力（μ）"],
            ["拥堵阈值", "2 500 人", "需启动疏导预案的队列规模（L）"],
            ["风险个体占比", "50 / 60 000", "低基率报警场景 P(T)"],
            ["灵敏度 / 误报率", "95% / 0.5%", "识别系统参数 P(B|T) / P(B|T̄)"],
            ["巡逻小组 / 热点偏向", "6 组 / 55%", "随机游走搜索能力"],
        ],
    )

    add_heading(doc, "3.3 叙事版交互作品（interactive/index.html）", 2)
    add_paragraph(
        doc,
        "叙事版交互作品采用纯 HTML + CSS + JavaScript 开发，无需安装任何依赖，"
        "直接在浏览器中打开 interactive/index.html 即可使用。"
        "页面结构采用左右双栏布局：左侧叙事面板以分幕引导的方式逐步呈现概率模型的现实含义，"
        "右侧 Canvas 区域动态绘制到达曲线、队列演化轨迹、巡逻路径和方案对比柱图。"
        "顶部六张指标卡实时显示当前参数下的关键数值，包括最大队列、拥堵概率、平均等待时间、"
        "报警后验概率、巡逻首达中位数和误报核查警员小时。"
        "用户通过底部滑杆调整参数后，系统立即在浏览器内重新运行蒙特卡洛模拟，"
        "使'参数改变-风险响应-策略解释'的因果链路直观可见。",
    )

    add_heading(doc, "3.4 数据仪表板版（webapp/）", 2)
    add_paragraph(
        doc,
        "数据仪表板版采用 React + Vite 技术栈开发，提供更系统化的多场景比较功能。"
        "仪表板将五个分析模块组织为可切换的选项卡，并以深色玻璃拟态风格呈现，"
        "配合渐变动效和微交互提升专业感。"
        "与叙事版互补：叙事版侧重引导用户逐步理解每个概率概念，"
        "仪表板版则适合在同一屏幕内对不同方案的多项指标进行横向比较。"
        "运行方式：在 webapp/ 目录下执行 npm install 后运行 npm run dev，"
        "浏览器访问本地开发服务器即可使用完整功能。",
    )
    add_paragraph(
        doc,
        "两套交互系统在数学模型和参数定义上完全对齐，所有随机种子均固定为 202405，"
        "保证相同参数设置下两版的数值结论一致，方便对照验证。",
    )

    # ── 第四章 结果 ─────────────────────────────────────────────────────────
    add_heading(doc, "第四章 模拟结果与概率解释", 1)
    add_heading(doc, "4.1 基准方案的单次路径与尾部风险", 2)
    doc.add_picture(str(figures / "arrival_queue_story.png"), width=Inches(5.9))
    add_caption(doc, "图 1  非齐次泊松到达、平均服务能力与队列演化（种子 202405）")
    add_paragraph(
        doc,
        f"图 1 展示了基准参数下一条固定随机路径（种子 202405）的演化过程。"
        f"在该路径中，最大队列为 {metrics['single_queue']['max_queue']:.0f} 人，"
        f"加权平均等待时间约 {metrics['single_queue']['avg_wait']:.2f} 分钟，"
        f"全程未触发 2500 人拥堵阈值。"
        f"然而这一'平静'的单次结果具有相当大的误导性："
        f"对 300 次独立蒙特卡洛模拟进行统计后，"
        f"最大队列超过阈值的风险概率为 {pct(metrics['queue']['risk_probability'])}，"
        f"平均最大队列达到 {metrics['queue']['max_queue_mean']:.0f} 人，"
        f"90 分位最大队列更高达 {metrics['queue']['max_queue_p90']:.0f} 人。"
        f"单次仿真与蒙特卡洛估计之间的显著差距说明：安防设计必须以尾部概率而非期望值作为决策依据，"
        f"否则极易低估真实风险。",
    )

    add_heading(doc, "4.2 通道数与到达集中度的敏感性分析", 2)
    doc.add_picture(str(figures / "queue_risk_heatmap.png"), width=Inches(5.9))
    add_caption(doc, "图 2  不同通道数与到达集中度组合下的拥堵概率热力图")
    add_paragraph(
        doc,
        "图 2 以热力图形式展示了通道数（34～58 个）与到达集中度（0.80～1.80）"
        "两个参数正交扫描时的拥堵风险概率。"
        "结果表明，拥堵风险对到达集中度的敏感程度显著高于对通道数的敏感程度："
        "当集中度从 0.80 升至 1.60 时，即使通道数保持在较高水平（50 个），"
        "拥堵风险仍可从接近 0% 升至 40% 以上。"
        "这一规律揭示了'错峰入场'的优先级应高于'单纯扩充通道'："
        "通过票面分时、预约入场等方式压低集中度，"
        "可以在不增加任何硬件成本的前提下大幅降低拥堵风险。"
        "两项措施配合使用时（组合优化方案），效果最为显著。",
    )

    add_heading(doc, "4.3 报警系统中的低基率效应", 2)
    doc.add_picture(str(figures / "bayes_alert_tradeoff.png"), width=Inches(5.9))
    add_caption(doc, "图 3  误报率对报警后验概率与警力核查成本的双轴影响")
    add_paragraph(
        doc,
        f"图 3 以双纵轴方式展示了误报率在 0.05%～2% 区间内变化时，"
        f"报警后验概率 P(T|B) 与警员核查小时数的同步变化趋势。"
        f"基准参数下（误报率 0.5%），预期真实报警 {metrics['recognition']['true_alerts']:.1f} 次，"
        f"误报约 {metrics['recognition']['false_alerts']:.1f} 次，"
        f"后验概率仅 {pct(metrics['recognition']['posterior_true_given_alert'])}，"
        f"误报核查耗费约 {metrics['recognition']['review_police_hours']:.0f} 警员小时。"
        f"曲线显示，当误报率从 0.5% 降至 0.1% 时，后验概率可提升至 40% 以上，"
        f"核查成本同步下降约 80%。"
        f"这一结果提示，在系统部署阶段提高识别精度（降低误报率）比事后增加人工核查力量更具成本效益。",
    )

    add_heading(doc, "4.4 巡逻偏向对首达时的影响", 2)
    doc.add_picture(str(figures / "patrol_hit_time.png"), width=Inches(5.9))
    add_caption(doc, "图 4  近随机巡逻与热点偏向巡逻的首达时分布对比（300 次模拟）")
    add_paragraph(
        doc,
        f"图 4 对比了 6 组巡逻队在近随机（偏向概率 10%）与热点偏向（偏向概率 55%）"
        f"两种策略下的首达时分布（各 300 次模拟）。"
        f"近随机巡逻的首达时中位数约为 {metrics['patrol_random']['median_hit_time']:.0f} 分钟，"
        f"分布较为分散；热点偏向巡逻的中位数降至约 {metrics['patrol_smart']['median_hit_time']:.0f} 分钟，"
        f"90 分位约为 {metrics['patrol_smart']['p90_hit_time']:.0f} 分钟，"
        f"分布明显向左收紧。"
        f"从安防角度看，中位数代表'多数情况下的发现速度'，"
        f"而 90 分位代表'极端情况下最迟发现时间'。"
        f"热点偏向在两项指标上均有明显优势，说明在已知拥堵热区、入口密度和报警位置的前提下，"
        f"有限的信息引导可以将搜索效率大幅提升，而无需增加巡逻队数量。",
    )

    # ── 第五章 策略建议 ─────────────────────────────────────────────────────
    add_heading(doc, "第五章 安防策略建议与方案比较", 1)
    add_paragraph(
        doc,
        "基于前述模拟结果，本章对四种典型安防方案进行量化比较（均使用 300 次蒙特卡洛模拟，种子 202405）。"
        "各方案从通道数、到达集中度、拥堵概率、平均最大队列、平均等待时间、"
        "巡逻首达中位数、报警后验概率和误报核查小时数八个维度进行评估，结果见下表。",
        first_line=True,
    )
    add_table(
        doc,
        ["方案", "通道数", "集中度", "拥堵概率", "平均最大队列\n(人)", "平均等待\n(分钟)", "巡逻首达\n中位数(分)", "报警\n可信度", "误报核查\n(警员小时)"],
        case_rows,
    )
    add_paragraph(
        doc,
        "从表中可以得出以下结论：",
    )
    add_paragraph(
        doc,
        "①  错峰入场（集中度降至 0.82）：在不增加任何硬件投入的前提下，"
        "将拥堵概率从 26.7% 降至 0.0%，平均最大队列下降约 42%，平均等待时间减半。"
        "这是成本最低、效果最显著的单一改善措施，应作为首选策略。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "②  增开通道（36 → 44 个）：对门区排队问题最为直接，"
        "拥堵概率同样降至 0%，但平均最大队列降幅更为显著（降至约 40 人）。"
        "不足之处是此举对报警系统和场内巡逻均无帮助，且硬件成本较高，"
        "在场地条件受限时难以大量扩充。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "③  组合优化（错峰 + 扩通道 + 降误报 + 增强偏向）：四个维度同步改善，"
        "不仅完全消除拥堵风险，还将巡逻首达时间从 5 分钟压缩至 3 分钟，"
        "报警后验概率从 13.7% 提升至 26.5%，误报核查成本下降超过 55%，"
        "是综合效果最优的方案。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "实践部署建议：先以预约入场和票面分时方式降低到达集中度（错峰），"
        "再保留适量通道冗余以应对随机波动，最后通过提升识别精度和热点引导巡逻减少警力浪费。"
        "三层措施分层推进，既可控制成本，也能覆盖从入口拥堵到场内异常的全链路安防需求。",
        first_line=True,
    )

    # ── 第六章 反思 ─────────────────────────────────────────────────────────
    add_heading(doc, "第六章 学习反思报告", 1)
    add_paragraph(
        doc,
        "本次作业最大的收获，在于将课本中的随机变量和随机过程真正放进一个有时间维度、"
        "有空间结构、有决策压力的现实情境里。"
        "项目初期，我们倾向于用平均到达率乘以时长来估算队列，"
        "但在推导排队递推公式后意识到：当每分钟到达量的波动被忽略时，"
        "期望值会系统性地低估最坏情况。"
        "蒙特卡洛的引入使'风险概率'从一个抽象概念变成了可以直接估计的数字，"
        "也让我们第一次直观感受到'期望值不等于风险'这一判断在实践中的重量。",
    )
    add_paragraph(
        doc,
        "贝叶斯后验概率部分是本项目的另一个认知突破。"
        "我们起初认为'灵敏度 95%'意味着报警大部分为真，"
        "计算后才发现在 0.083% 的基率下后验概率仅为 13.7%。"
        "这个反直觉的结论促使我们把误报率曲线（图 3）专门拉出来分析，"
        "最终得到'降低误报率的边际收益远大于提高灵敏度'的实用结论，"
        "并将其纳入策略建议。",
    )
    add_paragraph(
        doc,
        "在交互实现层面，我们分别开发了叙事版（纯前端，面向逐步引导）"
        "和仪表板版（React，面向多方案对比）两套系统，"
        "体会到'展示对象不同，交互设计就应当不同'。"
        "叙事版更注重把概率逻辑转化为文字与动画的配合，"
        "仪表板版则侧重让数据可以在同一屏幕内横向对齐比较。"
        "两套系统的并行开发还促使我们统一了随机种子、参数命名和指标定义，"
        "避免了数据口径不一致导致的结论冲突。",
    )
    add_paragraph(
        doc,
        "协作分工方面，模型推导、代码实现、图表解释和叙事结构之间需要反复对齐。"
        "每次参数调整都要同步更新仿真脚本、交互页面和报告文字，"
        "这促使我们建立了以 simulation.py 为唯一数据源、build_report.py 自动生成报告的工作流，"
        "有效避免了手动复制数据时的笔误。"
        "通过这次项目，我们更直观地感受到概率论不仅是计算工具，"
        "也是一种帮助公共管理进行不确定性决策的语言。",
    )

    # ── 参考文献 ─────────────────────────────────────────────────────
    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Ross, S. M. (2019). Introduction to Probability Models (12th ed.). Academic Press.",
        "[2] Lewis, P. A. W., & Shedler, G. S. (1979). Simulation of nonhomogeneous Poisson processes by thinning. Naval Research Logistics Quarterly, 26(3), 403-413.",
        "[3] Gilliam, R. R. (1979). An Application of Queueing Theory to Airport Passenger Security Screening. Interfaces, 9(4), 117-123.",
        "[4] Zhang, Z. G., Luh, H. P., & Wang, C. H. (2011). Modeling Security-Check Queues. Management Science, 57(11), 1979-1995.",
        "[5] Axelsson, S. (2000). The base-rate fallacy and the difficulty of intrusion detection. ACM Transactions on Information and System Security (TISSEC), 3(3), 186-205.",
        "[6] Lovász, L. (1993). Walks on graphs: A survey. Combinatorics, Paul erdos is eighty, 2(1), 1-46.",
    ]
    for ref in references:
        add_paragraph(doc, ref, first_line=False)

        # ── 附录 ─────────────────────────────────────────────────────────────────
    add_heading(doc, "附录（代码与运行说明）", 1)
    add_paragraph(
        doc,
        "完整源代码已随报告分别放入 source、interactive 和 webapp 文件夹。"
        "各套系统的运行方式如下：",
    )
    add_paragraph(doc, "（1）叙事版交互作品：直接在浏览器中打开 interactive/index.html 即可，无需安装任何依赖。", first_line=True)
    add_paragraph(doc, "（2）数据仪表板版：在 webapp/ 目录下执行 npm install，然后运行 npm run dev，浏览器访问命令行提示的本地地址即可。", first_line=True)
    add_paragraph(doc, "（3）复现实验图表：在命令行执行 py -X utf8 source/simulation.py --output report/figures，将重新生成全部图表和 metrics.json。", first_line=True)
    add_paragraph(doc, "（4）重新生成 Word 报告：安装 python-docx 后在 scripts/ 目录下执行 py build_report.py，将自动读取最新 metrics.json 并输出 report/midterm_report.docx。", first_line=True)
    add_paragraph(doc, "主要文件一览：", first_line=False)
    add_table(
        doc,
        ["文件 / 目录", "作用"],
        [
            ["interactive/index.html", "叙事版交互作品页面（纯前端，直接打开）"],
            ["interactive/styles.css", "叙事版样式与响应式布局"],
            ["interactive/app.js", "叙事版：浏览器内蒙特卡洛模拟与 Canvas 绘图"],
            ["webapp/src/", "数据仪表板版 React 源码（Vite 构建）"],
            ["source/simulation.py", "可复现仿真脚本：生成图表与 metrics.json"],
            ["scripts/build_report.py", "自动读取 metrics.json 并生成本 Word 报告"],
            ["report/figures/*.png", "报告中使用的四张可视化图表"],
            ["report/figures/metrics.json", "仿真输出的所有关键指标（数据唯一来源）"],
        ],
    )
    add_paragraph(
        doc,
        "核心队列递推伪代码（Python 风格）：",
        first_line=False,
    )
    add_formula(doc, "for t in range(minutes):")
    add_formula(doc, "    arrivals  = Poisson(lambda[t])")
    add_formula(doc, "    service   = Poisson(gates × service_rate)")
    add_formula(doc, "    queue[t]  = max(0, queue[t-1] + arrivals − service)")
    add_paragraph(
        doc,
        "重复上述过程 300 次并统计 max(queue) > threshold 的比例，"
        "即可可靠估计 P(max Q(t) > L)（尾部风险概率）。",
        first_line=False,
    )

    # ── 表格标题行底纹 ───────────────────────────────────────────────────────
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if row_idx == 0:
                    set_cell_shading(cell, "E8F3EF")

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
